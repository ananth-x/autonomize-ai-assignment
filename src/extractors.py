"""Document extraction module.

Handles extracting text content from various form formats:
- PDF (text-based and scanned/image-based)
- Images (PNG, JPG, TIFF)
- DOCX (Word documents)
- Plain text / CSV
"""

import os
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import fitz  # PyMuPDF
from PIL import Image

try:
    import pytesseract
except ImportError:
    pytesseract = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None


class FormExtractor:
    """Extracts text content from various document formats."""

    SUPPORTED_EXTENSIONS = {".pdf", ".png", ".jpg", ".jpeg", ".tiff", ".tif", ".docx", ".txt", ".csv"}

    def extract(self, file_path: str) -> dict:
        """Extract text and metadata from a form file.

        Args:
            file_path: Path to the form file.

        Returns:
            Dictionary with keys:
                - text: Extracted text content
                - metadata: Dict with source file info
                - pages: Number of pages (if applicable)
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = path.suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file format: {ext}. "
                f"Supported: {', '.join(sorted(self.SUPPORTED_EXTENSIONS))}"
            )

        metadata = {
            "source": str(path.absolute()),
            "filename": path.name,
            "extension": ext,
        }

        if ext == ".pdf":
            text, pages = self._extract_pdf(file_path)
            metadata["pages"] = pages
        elif ext in {".png", ".jpg", ".jpeg", ".tiff", ".tif"}:
            text = self._extract_image(file_path)
            metadata["pages"] = 1
        elif ext == ".docx":
            text = self._extract_docx(file_path)
            metadata["pages"] = 1
        elif ext in {".txt", ".csv"}:
            text = self._extract_text(file_path)
            metadata["pages"] = 1
        else:
            raise ValueError(f"No extractor for: {ext}")

        return {"text": text.strip(), "metadata": metadata}

    def _extract_pdf(self, file_path: str) -> Tuple[str, int]:
        """Extract text from PDF. Falls back to OCR for scanned pages."""
        doc = fitz.open(file_path)
        pages = len(doc)
        text_parts = []

        for page_num in range(pages):
            page = doc[page_num]
            text = page.get_text()

            # If page has very little text, try OCR
            if len(text.strip()) < 50 and pytesseract:
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                ocr_text = pytesseract.image_to_string(img)
                if len(ocr_text.strip()) > len(text.strip()):
                    text = ocr_text

            text_parts.append(f"[Page {page_num + 1}]\n{text}")

        doc.close()
        return "\n\n".join(text_parts), pages

    def _extract_image(self, file_path: str) -> str:
        """Extract text from image using OCR."""
        if pytesseract is None:
            raise ImportError(
                "pytesseract is required for image extraction. "
                "Install it with: pip install pytesseract"
            )
        img = Image.open(file_path)
        return pytesseract.image_to_string(img)

    def _extract_docx(self, file_path: str) -> str:
        """Extract text from Word document."""
        if DocxDocument is None:
            raise ImportError(
                "python-docx is required for DOCX extraction. "
                "Install it with: pip install python-docx"
            )
        doc = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        return "\n".join(paragraphs)

    def _extract_text(self, file_path: str) -> str:
        """Extract content from plain text or CSV files."""
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()


class StructuredFieldExtractor:
    """Extracts structured key-value fields from form text.

    Uses heuristics to identify common form patterns like:
    - Label: Value
    - Label _____ Value
    - Checkbox fields
    """

    def extract_fields(self, text: str) -> Dict[str, str]:
        """Extract key-value pairs from form text using pattern matching."""
        fields = {}
        lines = text.split("\n")

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Pattern: "Key: Value" or "Key : Value"
            if ":" in line:
                parts = line.split(":", 1)
                key = parts[0].strip()
                value = parts[1].strip()
                # Only treat as field if key is reasonably short (label-like)
                if key and value and len(key) < 60:
                    fields[key] = value

            # Pattern: "Key _____ Value" (underline separator)
            elif "___" in line:
                parts = line.split("___")
                key = parts[0].strip()
                value = parts[-1].strip()
                if key and value:
                    fields[key] = value

            # Pattern: "[x] Label" or "[ ] Label" (checkboxes)
            elif line.startswith("[") and "]" in line[:4]:
                checked = "x" in line[:4].lower() or "✓" in line[:4]
                label = line[line.index("]") + 1:].strip()
                if label:
                    fields[label] = "Yes" if checked else "No"

        return fields
